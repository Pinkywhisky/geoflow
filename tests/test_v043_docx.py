from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.documents import generate_copropriete_draft
from app.documents.generator import CONTENT_WIDTH_DXA
from app.domain import (
    Dossier,
    Millieme,
    Parcelle,
    StatutValidationMillieme,
)
from app.workflow import update_dossier_details
from tools.build_copropriete_template import build_template


FIXTURE = Path(__file__).parent / "fixtures" / "copropriete_complete.json"


def complete_business_dossier() -> Dossier:
    dossier = Dossier.model_validate_json(FIXTURE.read_text(encoding="utf-8"))
    update_dossier_details(
        dossier,
        numero="12",
        voie="rue des Écoles",
        complement="Bâtiment B",
        code_postal="75005",
        commune="Paris",
        departement="Paris",
        date_plan="2026-08-20",
    )
    dossier.references_cadastrales = [
        Parcelle(commune="Paris", section="AO", numero="163"),
        Parcelle(commune="Paris", section="AO", numero="164"),
    ]
    dossier.lots[0].usage = "Appartement"
    dossier.lots[0].designation = (
        "Appartement comprenant entrée, séjour, cuisine et chambre."
    )
    dossier.lots[1].usage = "Bureau"
    dossier.lots[1].designation = "Bureau situé au premier étage."
    dossier.milliemes = [
        Millieme(
            lot_id="lot-1",
            valeur=400,
            base=1000,
            statut_validation=StatutValidationMillieme.VALIDE,
        ),
        Millieme(
            lot_id="lot-2",
            valeur=600,
            base=1000,
            statut_validation=StatutValidationMillieme.VALIDE,
        ),
    ]
    dossier.grille_milliemes_complete = True
    return dossier


def text_content(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def test_generated_and_rebuilt_template_use_required_margins(tmp_path: Path) -> None:
    generated = generate_copropriete_draft(complete_business_dossier())
    document = Document(BytesIO(generated.content))
    template_path = tmp_path / "template.docx"
    build_template(template_path)
    template = Document(template_path)

    for candidate in (document, template):
        section = candidate.sections[0]
        assert section.left_margin.cm == pytest.approx(5, abs=0.001)
        assert section.right_margin.cm == pytest.approx(2, abs=0.001)
        assert section.top_margin.cm == pytest.approx(2, abs=0.001)
        assert section.bottom_margin.cm == pytest.approx(2, abs=0.001)


def test_generated_docx_defines_even_page_header_and_footer() -> None:
    generated = generate_copropriete_draft(complete_business_dossier())
    document = Document(BytesIO(generated.content))
    section = document.sections[0]

    assert document.settings.odd_and_even_pages_header_footer
    assert "GeoFlow" in section.even_page_header.paragraphs[0].text
    assert "BROUILLON" in section.even_page_footer.paragraphs[0].text


def test_docx_contains_completed_business_data() -> None:
    generated = generate_copropriete_draft(complete_business_dossier())
    text = text_content(Document(BytesIO(generated.content)))

    assert "12 rue des Écoles" in text
    assert "75005 Paris" in text
    assert "Paris — section AO n° 163" in text
    assert "Appartement" in text
    assert "Appartement comprenant entrée, séjour" in text
    assert "400 / 1000" in text
    assert "Total saisi : 1000 / 1000" in text


def test_docx_uses_business_zone_names_and_hides_internal_ids() -> None:
    dossier = complete_business_dossier()
    generated = generate_copropriete_draft(dossier)
    text = text_content(Document(BytesIO(generated.content)))

    assert "Zone 1" in text
    assert "Zone 2" in text
    assert "Partie commune 1" in text
    for zone in dossier.zones:
        assert zone.id not in text
        assert zone.geometrie_source.handle_dxf not in text


def test_all_docx_tables_fit_the_reduced_usable_width() -> None:
    generated = generate_copropriete_draft(complete_business_dossier())
    document = Document(BytesIO(generated.content))

    for table in document.tables:
        widths = [
            int(column.get(qn("w:w")))
            for column in table._tbl.tblGrid.gridCol_lst
        ]
        assert sum(widths) == CONTENT_WIDTH_DXA
