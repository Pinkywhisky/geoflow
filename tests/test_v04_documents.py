from __future__ import annotations

import hashlib
import re
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from app.documents import (
    TEMPLATE_ID,
    TEMPLATE_VERSION,
    GenerationBlockedError,
    TemplateCorruptError,
    TemplateMissingError,
    assess_generation_readiness,
    generate_copropriete_draft,
)
from app.documents.generator import canonical_snapshot, safe_document_filename
from app.domain import Dossier, Millieme
from app.storage import DossierNotFoundError, JsonDossierRepository
from tools.build_copropriete_template import build_template


FIXTURE = Path(__file__).parent / "fixtures" / "copropriete_complete.json"
GENERATED_AT = datetime(2026, 7, 1, 12, 30, tzinfo=timezone.utc)


@pytest.fixture
def complete_dossier() -> Dossier:
    return Dossier.model_validate_json(FIXTURE.read_text(encoding="utf-8"))


@pytest.fixture
def generated(complete_dossier: Dossier):
    return generate_copropriete_draft(
        complete_dossier, generated_at=GENERATED_AT
    )


def document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def open_generated(generated) -> Document:
    return Document(BytesIO(generated.content))


def test_complete_fixture_is_ready(complete_dossier: Dossier) -> None:
    readiness = assess_generation_readiness(complete_dossier)
    assert readiness.can_generate
    assert readiness.blockers == ()
    assert readiness.retained_zone_count == 4
    assert readiness.corrected_surface_count == 1
    assert readiness.retained_surface_m2 == pytest.approx(182)


def test_generated_bytes_are_a_readable_docx(generated) -> None:
    document = open_generated(generated)
    assert generated.content.startswith(b"PK")
    assert len(document.sections) == 1
    assert len(document.tables) >= 6


def test_cover_contains_reference_address_and_draft_status(generated) -> None:
    text = document_text(open_generated(generated))
    assert "DEMO-2026-001" in text
    assert "10 avenue des Tilleuls" in text
    assert "BROUILLON À VALIDER" in text
    assert "ne constitue ni un acte authentique" in text


def test_all_lots_are_generated_in_natural_order(generated) -> None:
    text = document_text(open_generated(generated))
    assert "Lot 1" in text
    assert "Lot 2" in text
    assert text.index("Lot 1") < text.index("Lot 2")


def test_multi_zone_and_multi_level_lot_is_described(generated) -> None:
    text = document_text(open_generated(generated))
    assert "zone-lot-1-principale" in text
    assert "zone-lot-1-cave" in text
    assert "RdC, Sous-sol" in text
    assert "Appartement et cave" in text


def test_dynamic_recap_contains_retained_lot_totals(generated) -> None:
    document = open_generated(generated)
    recap = next(
        table
        for table in document.tables
        if table.cell(0, 0).text == "Lot"
        and table.cell(0, 4).text == "Surface retenue"
    )
    rows = [[cell.text for cell in row.cells] for row in recap.rows]
    assert any(row[0] == "1" and "100,00 m²" in row for row in rows)
    assert any(row[0] == "2" and "70,00 m²" in row for row in rows)
    assert "170,00 m²" in rows[-1]


def test_geometric_surface_correction_and_justification_are_visible(
    generated,
) -> None:
    text = document_text(open_generated(generated))
    assert "72,00 m²" in text
    assert "70,00 m² (-2.00 m²)" in text
    assert "Déduction d'une trémie mesurée et validée." in text


def test_common_zone_and_optional_right_are_visible(generated) -> None:
    text = document_text(open_generated(generated))
    assert "Parties communes identifiées" in text
    assert "zone-commune" in text
    assert "Droit d'usage privatif" in text
    assert "À confirmer" in text


def test_supplied_milliemes_are_displayed_without_calculation(
    complete_dossier: Dossier,
) -> None:
    dossier = complete_dossier.model_copy(deep=True)
    dossier.milliemes = [
        Millieme(lot_id="lot-1", valeur=250, base=1000, valide=False)
    ]
    generated = generate_copropriete_draft(
        dossier, generated_at=GENERATED_AT
    )
    text = document_text(open_generated(generated))
    assert "Millièmes fournis" in text
    assert "sans calcul ni déduction automatique" in text
    assert "250" in text
    assert "1000" in text
    assert "À valider" in text


def test_remaining_validation_warnings_are_visible(generated) -> None:
    text = document_text(open_generated(generated))
    assert "Points restant à valider" in text
    assert "millièmes ne sont pas renseignés" in text
    assert "servitudes restent à confirmer" in text
    assert "clauses juridiques restent à compléter" in text
    assert "correction(s) manuelle(s)" in text


def test_unknown_optional_identification_is_omitted(
    complete_dossier: Dossier,
) -> None:
    dossier = complete_dossier.model_copy(deep=True)
    dossier.adresse = None
    dossier.commune = None
    dossier.references_cadastrales = []
    generated = generate_copropriete_draft(
        dossier, generated_at=GENERATED_AT
    )
    text = document_text(open_generated(generated))
    assert "10 avenue des Tilleuls" not in text
    assert "Section AB n° 101" not in text
    assert "Les références cadastrales restent à renseigner." in text


def test_docx_metadata_is_clean_and_generic(generated) -> None:
    core = open_generated(generated).core_properties
    assert core.author == "GeoFlow"
    assert core.last_modified_by == "GeoFlow"
    assert core.subject == "Copropriété"
    assert core.title == "GeoFlow — DEMO-2026-001 — Brouillon"


def test_filename_is_neutral_and_cannot_traverse() -> None:
    filename = safe_document_filename("../../Dossier Été 2026")
    assert filename == "geoflow_dossier_ete_2026_copropriete_brouillon.docx"
    assert "/" not in filename
    assert "\\" not in filename
    assert ".." not in filename


def test_generation_records_template_and_snapshot_hash(
    complete_dossier: Dossier, generated
) -> None:
    assert generated.generation.template_id == TEMPLATE_ID
    assert generated.generation.template_version == TEMPLATE_VERSION
    assert generated.generation.statut == "brouillon"
    assert generated.generation.date_generation == GENERATED_AT
    assert generated.generation.sha256_snapshot == hashlib.sha256(
        generated.snapshot
    ).hexdigest()
    assert generated.snapshot == canonical_snapshot(complete_dossier)
    assert b'"generations":[]' in generated.snapshot


def test_generator_does_not_require_the_source_plan(
    complete_dossier: Dossier,
) -> None:
    complete_dossier.plan_importe.nom_fichier_original = (
        "source_absente_mais_tracee.dxf"
    )
    generated = generate_copropriete_draft(
        complete_dossier, generated_at=GENERATED_AT
    )
    assert "source_absente_mais_tracee.dxf" in document_text(
        open_generated(generated)
    )


def test_package_has_no_comments_revisions_external_links_or_local_paths(
    generated,
) -> None:
    with ZipFile(BytesIO(generated.content)) as archive:
        names = archive.namelist()
        assert not any(name.startswith("word/comments") for name in names)
        assert "docProps/custom.xml" not in names
        xml = b"".join(
            archive.read(name)
            for name in names
            if name.endswith((".xml", ".rels"))
        )
    assert re.search(br"<w:(?:ins|del)(?:[ >])", xml) is None
    assert b'TargetMode="External"' not in xml
    assert b"file:///" not in xml
    assert b"C:\\Users\\" not in xml
    assert b"/home/" not in xml


def test_missing_template_has_a_specific_error(
    complete_dossier: Dossier, tmp_path: Path
) -> None:
    with pytest.raises(TemplateMissingError, match="introuvable"):
        generate_copropriete_draft(
            complete_dossier, template_path=tmp_path / "missing.docx"
        )


def test_corrupt_template_has_a_specific_error(
    complete_dossier: Dossier, tmp_path: Path
) -> None:
    template = tmp_path / "corrupt.docx"
    template.write_bytes(b"not a zip")
    with pytest.raises(TemplateCorruptError, match="illisible"):
        generate_copropriete_draft(
            complete_dossier, template_path=template
        )


@pytest.mark.parametrize(
    "mutation, expected",
    [
        (
            lambda dossier: setattr(
                dossier.plan_importe, "unite_confirmee", False
            ),
            "unité du plan",
        ),
        (
            lambda dossier: setattr(
                dossier.planches[0], "statut", "candidate"
            ),
            "Aucune planche",
        ),
        (
            lambda dossier: setattr(
                dossier.zones[0], "lot_id", None
            ),
            "aucun lot",
        ),
    ],
)
def test_generation_is_blocked_on_invalid_canonical_state(
    complete_dossier: Dossier, mutation, expected: str
) -> None:
    dossier = complete_dossier.model_copy(deep=True)
    mutation(dossier)
    readiness = assess_generation_readiness(dossier)
    assert not readiness.can_generate
    assert any(expected in item for item in readiness.blockers)
    with pytest.raises(GenerationBlockedError):
        generate_copropriete_draft(dossier)


def test_generation_artifacts_are_isolated_and_path_safe(
    tmp_path: Path, complete_dossier: Dossier, generated
) -> None:
    repository = JsonDossierRepository(tmp_path / "data")
    generation = generated.generation
    repository.save_generation_artifacts(
        complete_dossier.id,
        generation.id,
        generation.nom_fichier,
        generated.content,
        generated.snapshot,
    )
    document_path = repository.generated_document_path(
        complete_dossier.id, generation.id, generation.nom_fichier
    )
    assert document_path.read_bytes() == generated.content
    assert (document_path.parent / "dossier_snapshot.json").read_bytes() == (
        generated.snapshot
    )
    with pytest.raises(DossierNotFoundError):
        repository.generated_document_path(
            complete_dossier.id, "../private", generation.nom_fichier
        )
    with pytest.raises(ValueError):
        repository.save_generation_artifacts(
            complete_dossier.id,
            "other",
            "../escape.docx",
            generated.content,
            generated.snapshot,
        )


def test_template_builder_is_byte_reproducible(tmp_path: Path) -> None:
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    build_template(first)
    build_template(second)
    assert first.read_bytes() == second.read_bytes()
    assert Document(first).core_properties.author == "GeoFlow"
