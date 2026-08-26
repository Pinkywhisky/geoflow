"""Generate the first versioned copropriete draft from canonical JSON data."""

from __future__ import annotations

import hashlib
import json
import math
import re
import unicodedata
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from uuid import uuid4
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips

from app.domain import (
    CategorieZone,
    Dossier,
    Generation,
    StatutValidationMillieme,
    StatutValidationJuridique,
)
from app.domain.models import StatutRevue

from .context import GenerationReadiness, assess_generation_readiness


TEMPLATE_ID = "copropriete_draft_v1"
TEMPLATE_VERSION = "1.1"
DEFAULT_TEMPLATE_PATH = (
    Path(__file__).resolve().parent / "templates" / f"{TEMPLATE_ID}.docx"
)
DOCUMENT_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
# A4 usable width with 5 cm left and 2 cm right margins is 7,937 DXA.
# Tables keep a 120 DXA indent and therefore use the remaining 7,817 DXA.
CONTENT_WIDTH_DXA = 7817


class DocumentGenerationError(RuntimeError):
    pass


class TemplateMissingError(DocumentGenerationError):
    pass


class TemplateCorruptError(DocumentGenerationError):
    pass


class GenerationBlockedError(DocumentGenerationError):
    def __init__(self, blockers: tuple[str, ...]) -> None:
        super().__init__("Le dossier ne permet pas la génération du document.")
        self.blockers = blockers


@dataclass(frozen=True)
class GeneratedDocument:
    content: bytes
    snapshot: bytes
    generation: Generation


def canonical_snapshot(dossier: Dossier) -> bytes:
    payload = dossier.model_dump(mode="json")
    return (
        json.dumps(
            payload,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
        + b"\n"
    )


def safe_document_filename(reference: str) -> str:
    normalized = unicodedata.normalize("NFKD", reference)
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii").lower()
    slug = re.sub(r"[^a-z0-9]+", "_", ascii_value).strip("_")[:60] or "dossier"
    return f"geoflow_{slug}_copropriete_brouillon.docx"


def generate_copropriete_draft(
    dossier: Dossier,
    *,
    template_path: str | Path = DEFAULT_TEMPLATE_PATH,
    generated_at: datetime | None = None,
) -> GeneratedDocument:
    """Create a DOCX using only the supplied canonical dossier snapshot."""

    readiness = assess_generation_readiness(dossier)
    if not readiness.can_generate:
        raise GenerationBlockedError(readiness.blockers)

    template = Path(template_path)
    if not template.is_file():
        raise TemplateMissingError(f"Template introuvable : {template.name}")
    try:
        document = Document(str(template))
    except (PackageNotFoundError, BadZipFile, KeyError, ValueError) as exc:
        raise TemplateCorruptError("Le template DOCX est illisible.") from exc

    _configure_document(document)
    _clear_body(document)
    when = generated_at or datetime.now(timezone.utc)
    if when.tzinfo is None:
        when = when.replace(tzinfo=timezone.utc)
    _set_metadata(document, dossier)
    _write_document(document, dossier, readiness, when)

    output = BytesIO()
    try:
        document.save(output)
    except (OSError, ValueError, KeyError) as exc:
        raise TemplateCorruptError(
            "Le template DOCX ne peut pas être enregistré."
        ) from exc
    content = output.getvalue()
    _assert_clean_package(content)

    snapshot = canonical_snapshot(dossier)
    generation = Generation(
        id=uuid4().hex,
        type_document="etat_descriptif_copropriete",
        template_id=TEMPLATE_ID,
        template_version=TEMPLATE_VERSION,
        date_generation=when,
        sha256_snapshot=hashlib.sha256(snapshot).hexdigest(),
        statut="brouillon",
        nom_fichier=safe_document_filename(dossier.reference),
        avertissements=list(readiness.warnings),
    )
    return GeneratedDocument(
        content=content,
        snapshot=snapshot,
        generation=generation,
    )


def _configure_document(document: Document) -> None:
    # Explicit even-page parts are required for reliable DOCX interoperability:
    # LibreOffice otherwise imports only the right-page header/footer and lets
    # body content use the header area on left (even) pages.
    document.settings.odd_and_even_pages_header_footer = True
    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(5)
        section.right_margin = Cm(2)
        section.header_distance = Cm(1.15)
        section.footer_distance = Cm(1.15)
        _configure_header(section.header)
        _configure_footer(section.footer)
        _configure_header(section.even_page_header)
        _configure_footer(section.even_page_footer)
    _configure_styles(document)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x24, 0x2D, 0x2A)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_settings = {
        "Heading 1": (16, RGBColor(0x2E, 0x74, 0xB5), 16, 8),
        "Heading 2": (13, RGBColor(0x2E, 0x74, 0xB5), 12, 6),
        "Heading 3": (12, RGBColor(0x1F, 0x4D, 0x78), 8, 4),
    }
    for name, (size, color, before, after) in heading_settings.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        # LibreOffice can move keep-with-next groups into the header area when
        # a group overflows. Normal pagination is more reliable for this draft.
        style.paragraph_format.keep_with_next = False

    custom_styles = {
        "GeoFlow Document Title": (26, True, RGBColor(0x12, 0x52, 0x43)),
        "GeoFlow Subtitle": (13, False, RGBColor(0x4D, 0x60, 0x59)),
        "GeoFlow Warning": (10, True, RGBColor(0x71, 0x50, 0x0F)),
        "GeoFlow Table Text": (9, False, RGBColor(0x24, 0x2D, 0x2A)),
        "GeoFlow Caption": (8, False, RGBColor(0x63, 0x71, 0x6C)),
    }
    for name, (size, bold, color) in custom_styles.items():
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_after = Pt(6)
    list_bullet = styles["List Bullet"]
    list_bullet.font.name = "Calibri"
    list_bullet.font.size = Pt(10)
    list_bullet.paragraph_format.left_indent = Inches(0.45)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.22)
    list_bullet.paragraph_format.space_after = Pt(8)



def _configure_header(header: object) -> None:
    paragraph = header.paragraphs[0]
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("GeoFlow  ·  Copropriété — Brouillon")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x63, 0x71, 0x6C)


def _configure_footer(footer: object) -> None:
    paragraph = footer.paragraphs[0]
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("BROUILLON À VALIDER  ·  Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x63, 0x71, 0x6C)
    _add_page_field(paragraph)


def _add_page_field(paragraph: object) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, value, end):
        run._r.append(element)


def _clear_paragraph(paragraph: object) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def _clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _set_metadata(document: Document, dossier: Dossier) -> None:
    core = document.core_properties
    core.title = f"GeoFlow — {dossier.reference} — Brouillon"
    core.subject = "Copropriété"
    core.author = "GeoFlow"
    core.last_modified_by = "GeoFlow"
    core.comments = "Document de travail généré depuis le dossier canonique GeoFlow."
    core.keywords = "GeoFlow, copropriété, brouillon"
    core.category = "Document de travail"


def _write_document(
    document: Document,
    dossier: Dossier,
    readiness: GenerationReadiness,
    generated_at: datetime,
) -> None:
    _write_cover(document, dossier, generated_at)
    page_start = document.add_paragraph()
    page_start.paragraph_format.page_break_before = True
    page_start.paragraph_format.space_after = Pt(0)
    # LibreOffice collapses a completely empty page-break paragraph when the
    # next block is a table, then lays that table above the configured top
    # margin. A one-point non-breaking-space anchor keeps the 2 cm margin in
    # both Word and the headless LibreOffice renderer used for visual QA.
    anchor = page_start.add_run("\u00a0")
    anchor.font.size = Pt(1)
    _write_draft_notice(document)
    _write_identification(document, dossier)
    _write_organization(document, dossier)
    _write_lots(document, dossier)
    _write_recap(document, dossier)
    _write_surface_control(document, dossier)
    _write_common_and_legal_elements(document, dossier)
    _write_millieme_section(document, dossier)
    _write_legal_placeholders(document)
    _write_remaining_points(document, readiness)
    paragraph = document.add_paragraph(style="GeoFlow Caption")
    paragraph.add_run(
        f"Template {TEMPLATE_ID} — version {TEMPLATE_VERSION} · "
        f"généré le {generated_at.astimezone(timezone.utc).strftime('%d/%m/%Y à %H:%M UTC')}."
    )


def _write_cover(
    document: Document, dossier: Dossier, generated_at: datetime
) -> None:
    brand = document.add_paragraph()
    brand.paragraph_format.space_before = Pt(42)
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run("GeoFlow")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x17, 0x6B, 0x55)

    title = document.add_paragraph(style="GeoFlow Document Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(42)
    title.add_run("ÉTAT DESCRIPTIF DE DIVISION\nET RÈGLEMENT DE COPROPRIÉTÉ")

    badge = document.add_paragraph(style="GeoFlow Warning")
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge.paragraph_format.space_before = Pt(24)
    badge.add_run("BROUILLON À VALIDER")

    details = [f"Référence : {dossier.reference}"]
    address = _formatted_address(dossier)
    cadastre = _formatted_cadastre(dossier)
    if address:
        details.append(f"Adresse : {address}")
    if dossier.commune:
        details.append(f"Commune : {dossier.commune}")
    if cadastre:
        details.append(f"Références cadastrales : {cadastre}")
    details.extend(
        [
            f"Date de génération : {generated_at.strftime('%d/%m/%Y')}",
            f"Template : {TEMPLATE_ID} — version {TEMPLATE_VERSION}",
        ]
    )
    subtitle = document.add_paragraph(style="GeoFlow Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(36)
    subtitle.add_run("\n".join(details))


def _write_draft_notice(document: Document) -> None:
    table = _add_table(document, [CONTENT_WIDTH_DXA], has_header=False)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "FFF5DC")
    paragraph = cell.paragraphs[0]
    paragraph.style = "GeoFlow Warning"
    paragraph.add_run(
        "Document de travail — Ce document est un brouillon généré à partir des "
        "données validées dans GeoFlow. Il ne constitue ni un acte authentique, "
        "ni une validation juridique, ni un règlement définitif."
    )


def _write_identification(document: Document, dossier: Dossier) -> None:
    document.add_heading("1. Identification de l’opération", level=1)
    rows: list[tuple[str, str]] = [
        ("Référence du dossier", dossier.reference),
        ("Adresse", _formatted_address(dossier) or "Non renseignée"),
        ("Commune", dossier.commune or "Non renseignée"),
        ("Département", dossier.departement or "Non renseigné"),
        (
            "Références cadastrales",
            _formatted_cadastre(dossier) or "Non renseignées",
        ),
        ("Date du plan / relevé", dossier.date_plan or "Non renseignée"),
        (
            "Plan source",
            dossier.plan_importe.nom_fichier_original
            if dossier.plan_importe
            else "Non renseigné",
        ),
    ]
    rows.extend(
        [
            ("Nombre de bâtiments", str(len(dossier.batiments))),
            ("Nombre de niveaux", str(len(dossier.niveaux))),
            ("Nombre de lots", str(len(dossier.lots))),
        ]
    )
    table = _add_table(
        document, [2200, CONTENT_WIDTH_DXA - 2200], has_header=False
    )
    for index, (label, value) in enumerate(rows):
        cells = table.rows[0].cells if index == 0 else table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = value
        _style_row(cells)


def _write_organization(document: Document, dossier: Dossier) -> None:
    document.add_heading("2. Organisation de l’ensemble immobilier", level=1)
    levels_by_building: dict[str, list[object]] = {}
    for level in dossier.niveaux:
        levels_by_building.setdefault(level.batiment_id, []).append(level)
    for building in dossier.batiments:
        title = f"Bâtiment {building.code}"
        if building.libelle:
            title += f" — {building.libelle}"
        document.add_heading(title, level=2)
        levels = levels_by_building.get(building.id, [])
        paragraph = document.add_paragraph()
        paragraph.add_run("Niveaux : ").bold = True
        paragraph.add_run(
            ", ".join(
                f"{level.code}{f' — {level.libelle}' if level.libelle else ''}"
                for level in levels
            )
            or "À renseigner"
        )


def _write_lots(document: Document, dossier: Dossier) -> None:
    document.add_heading("3. Description des lots", level=1)
    buildings = {item.id: item for item in dossier.batiments}
    levels = {item.id: item for item in dossier.niveaux}
    zones = {item.id: item for item in dossier.zones}
    zone_labels = _zone_labels(dossier)
    for lot in sorted(dossier.lots, key=lambda item: _natural_key(item.numero)):
        lot_heading = document.add_heading(f"Lot {lot.numero}", level=2)
        lot_heading.paragraph_format.keep_with_next = True
        facts = document.add_paragraph()
        facts.paragraph_format.keep_with_next = True
        facts.add_run("Bâtiment(s) : ").bold = True
        facts.add_run(
            ", ".join(
                buildings[item].code for item in lot.batiment_ids if item in buildings
            )
            or "Non renseigné"
        )
        facts.add_run("  ·  Niveau(x) : ").bold = True
        facts.add_run(
            ", ".join(
                levels[item].code for item in lot.niveau_ids if item in levels
            )
            or "Non renseigné"
        )
        facts.add_run("  ·  Usage : ").bold = True
        facts.add_run(lot.usage or "Non renseigné")
        designation = document.add_paragraph()
        designation.paragraph_format.keep_with_next = True
        designation.add_run("Désignation : ").bold = True
        designation.add_run(lot.designation or "Non renseignée")

        table = _add_table(
            document,
            [900, 2100, 1600, CONTENT_WIDTH_DXA - 4600],
            headers=(
                "Zone",
                "Catégorie",
                "Niveau",
                "Surface retenue",
            ),
        )
        rendered_zones = []
        for zone_id in lot.zone_ids:
            zone = zones.get(zone_id)
            if zone is None:
                continue
            rendered_zones.append(zone)
            values = (
                zone_labels[zone.id],
                zone.categorie.value.replace("_", " "),
                levels.get(zone.niveau_id).code
                if zone.niveau_id in levels
                else "—",
                _surface(zone.surface_retenue_m2),
            )
            _fill_row(table.add_row().cells, values)
        _keep_table_rows_together(table)
        total = sum(
            zones[zone_id].surface_retenue_m2
            for zone_id in lot.zone_ids
            if zone_id in zones
            and zones[zone_id].statut == StatutRevue.RETENUE
            and zones[zone_id].categorie != CategorieZone.EXCLUE
        )
        if len(rendered_zones) > 1:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.add_run(
                f"Surface retenue du lot : {_surface(total)}"
            ).bold = True


def _write_recap(document: Document, dossier: Dossier) -> None:
    document.add_heading("4. Tableau récapitulatif des lots", level=1)
    buildings = {item.id: item for item in dossier.batiments}
    levels = {item.id: item for item in dossier.niveaux}
    zones = {item.id: item for item in dossier.zones}
    milliemes = {item.lot_id: item for item in dossier.milliemes}
    if milliemes:
        widths = [650, 1150, 1050, 2250, 1317, 1400]
        headers = (
            "Lot",
            "Bât.",
            "Niveaux",
            "Usage / désignation",
            "Surface",
            "Millièmes",
        )
    else:
        widths = [700, 1150, 1100, 3200, 1667]
        headers = (
            "Lot",
            "Bât.",
            "Niveaux",
            "Usage / désignation",
            "Surface retenue",
        )
    table = _add_table(document, widths, headers=headers)
    grand_total = 0.0
    for lot in sorted(dossier.lots, key=lambda item: _natural_key(item.numero)):
        total = sum(
            zones[zone_id].surface_retenue_m2
            for zone_id in lot.zone_ids
            if zone_id in zones
            and zones[zone_id].statut == StatutRevue.RETENUE
            and zones[zone_id].categorie != CategorieZone.EXCLUE
        )
        grand_total += total
        description = " — ".join(
            value for value in (lot.usage, lot.designation) if value
        ) or "Non renseigné"
        values = [
            lot.numero,
            ", ".join(
                buildings[item].code
                for item in lot.batiment_ids
                if item in buildings
            ),
            ", ".join(
                levels[item].code
                for item in lot.niveau_ids
                if item in levels
            ),
            description,
            _surface(total),
        ]
        millieme = milliemes.get(lot.id)
        if milliemes:
            values.append(
                (
                    f"{millieme.valeur} / {millieme.base or 1000}"
                    if millieme
                    else "Non renseignés"
                )
            )
        _fill_row(
            table.add_row().cells,
            values,
        )
    cells = table.add_row().cells
    _merge_cells(cells, 0, 3)
    cells[0].text = "Total des surfaces privatives retenues"
    cells[0].paragraphs[0].runs[0].bold = True
    cells[4].text = _surface(grand_total)
    cells[4].paragraphs[0].runs[0].bold = True
    if milliemes:
        cells[5].text = f"{sum(item.valeur for item in milliemes.values())} / 1000"
        cells[5].paragraphs[0].runs[0].bold = True
    _style_row(cells)


def _write_surface_control(document: Document, dossier: Dossier) -> None:
    corrected = [
        zone
        for zone in dossier.zones
        if zone.statut == StatutRevue.RETENUE
        and not math.isclose(
            zone.surface_geometrique_m2,
            zone.surface_retenue_m2,
            abs_tol=1e-9,
        )
    ]
    document.add_heading("5. Contrôle des surfaces", level=1)
    if not corrected:
        document.add_paragraph(
            "Aucune correction manuelle de surface n'est enregistrée."
        )
        return
    lots = {item.id: item for item in dossier.lots}
    zone_labels = _zone_labels(dossier)
    for zone in corrected:
        lot = lots.get(zone.lot_id or "")
        difference = zone.surface_retenue_m2 - zone.surface_geometrique_m2
        document.add_heading(
            (
                f"Lot {lot.numero} — {zone_labels[zone.id]}"
                if lot
                else zone_labels[zone.id]
            ),
            level=3,
        )
        table = _add_table(
            document,
            [2100, CONTENT_WIDTH_DXA - 2100],
            has_header=False,
        )
        values = (
            ("Surface géométrique", _surface(zone.surface_geometrique_m2)),
            ("Surface retenue", _surface(zone.surface_retenue_m2)),
            ("Écart", f"{difference:+.2f} m²"),
            (
                "Justification",
                zone.decision_surface.justification or "Non renseignée",
            ),
        )
        for index, (label, value) in enumerate(values):
            cells = table.rows[0].cells if index == 0 else table.add_row().cells
            cells[0].text = label
            cells[0].paragraphs[0].runs[0].bold = True
            cells[1].text = value
            _style_row(cells)


def _write_common_and_legal_elements(
    document: Document, dossier: Dossier
) -> None:
    common_zones = [
        zone
        for zone in dossier.zones
        if zone.statut == StatutRevue.RETENUE
        and zone.categorie == CategorieZone.COMMUNE
    ]
    document.add_heading("6. Droits et servitudes éventuels", level=1)
    if not (common_zones or dossier.droits_particuliers or dossier.servitudes):
        document.add_paragraph(
            "Aucun droit particulier ni aucune servitude n'est renseigné."
        )
        return
    if common_zones:
        document.add_heading("Parties communes identifiées", level=2)
        table = _add_table(
            document,
            [1200, 1800, 1800, CONTENT_WIDTH_DXA - 4800],
            headers=("Zone", "Bâtiment", "Niveau", "Surface retenue"),
        )
        buildings = {item.id: item for item in dossier.batiments}
        levels = {item.id: item for item in dossier.niveaux}
        zone_labels = _zone_labels(dossier)
        for zone in common_zones:
            _fill_row(
                table.add_row().cells,
                (
                    zone_labels[zone.id],
                    buildings.get(zone.batiment_id).code
                    if zone.batiment_id in buildings
                    else "—",
                    levels.get(zone.niveau_id).code
                    if zone.niveau_id in levels
                    else "—",
                    _surface(zone.surface_retenue_m2),
                ),
            )
    if dossier.droits_particuliers:
        document.add_heading("Droits particuliers", level=2)
        lots = {item.id: item.numero for item in dossier.lots}
        for right in dossier.droits_particuliers:
            paragraph = document.add_paragraph()
            paragraph.add_run(right.description)
            if right.lot_ids:
                paragraph.add_run(
                    " — lot(s) "
                    + ", ".join(
                        lots.get(item, "lot inconnu") for item in right.lot_ids
                    )
                )
            if (
                right.statut_validation
                != StatutValidationJuridique.CONFIRME
            ):
                paragraph.add_run(" — À confirmer").bold = True
    if dossier.servitudes:
        document.add_heading("Servitudes", level=2)
        for servitude in dossier.servitudes:
            paragraph = document.add_paragraph()
            paragraph.add_run(servitude.description)
            if (
                servitude.statut_validation
                != StatutValidationJuridique.CONFIRME
            ):
                paragraph.add_run(" — À confirmer").bold = True
def _write_millieme_section(document: Document, dossier: Dossier) -> None:
    if not dossier.milliemes:
        return
    document.add_heading("7. Millièmes fournis", level=1)
    document.add_paragraph(
        "Valeurs reprises du dossier canonique sans calcul ni déduction "
        "automatique, et sans correction d'arrondi."
    )
    lots = {item.id: item.numero for item in dossier.lots}
    table = _add_table(
        document,
        [1500, 1500, 1300, CONTENT_WIDTH_DXA - 4300],
        headers=("Lot", "Valeur", "Base", "Statut"),
    )
    for millieme in dossier.milliemes:
        _fill_row(
            table.add_row().cells,
            (
                lots.get(millieme.lot_id, "Lot inconnu"),
                str(millieme.valeur),
                str(millieme.base) if millieme.base is not None else "1000",
                (
                    "Validé"
                    if millieme.statut_validation
                    == StatutValidationMillieme.VALIDE
                    else "À confirmer"
                ),
            ),
        )
    paragraph = document.add_paragraph()
    paragraph.add_run("Total saisi : ").bold = True
    paragraph.add_run(
        f"{sum(item.valeur for item in dossier.milliemes)} / 1000"
    )
    if dossier.grille_milliemes_complete:
        paragraph.add_run(" — grille déclarée complète").bold = True


def _write_legal_placeholders(document: Document) -> None:
    document.add_heading("8. Clauses juridiques / limites", level=1)
    document.add_paragraph(
        "La rédaction et la validation des clauses juridiques ne sont pas "
        "gérées dans cette version de GeoFlow. Le présent document reste un "
        "brouillon destiné à être contrôlé par les professionnels compétents."
    )
    document.add_paragraph(
        "Aucune qualification juridique n’est déduite automatiquement des "
        "géométries, annotations ou calques du plan source."
    )


def _write_remaining_points(
    document: Document, readiness: GenerationReadiness
) -> None:
    document.add_heading("9. Points restant à valider", level=1)
    if not readiness.user_actions:
        document.add_paragraph(
            "Aucune donnée métier éditable n'est signalée comme manquante."
        )
    for notice in readiness.user_actions:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(notice.message)
        run.bold = True
        run.font.color.rgb = RGBColor(0x71, 0x50, 0x0F)
    if readiness.product_limitations:
        document.add_heading("Limitations du produit", level=2)
        for notice in readiness.product_limitations:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(notice.message)
    for notice in readiness.information:
        paragraph = document.add_paragraph(style="GeoFlow Caption")
        paragraph.add_run(notice.message)


def _add_table(
    document: Document,
    widths: list[int],
    *,
    headers: tuple[str, ...] | None = None,
    has_header: bool = True,
):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(
            "La largeur du tableau doit correspondre à la zone utile A4."
        )
    table = document.add_table(rows=1, cols=len(widths))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    properties = table._tbl.tblPr
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = properties.find(qn("w:tblW"))
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        _set_row_widths(row.cells, widths)
    if headers is not None:
        _fill_row(table.rows[0].cells, headers)
        _format_header_row(table.rows[0])
    elif has_header:
        _format_header_row(table.rows[0])
    return table


def _set_row_widths(cells: object, widths: list[int]) -> None:
    for cell, width in zip(cells, widths, strict=True):
        cell.width = Twips(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        properties = cell._tc.get_or_add_tcPr()
        cell_width = properties.find(qn("w:tcW"))
        if cell_width is None:
            cell_width = OxmlElement("w:tcW")
            properties.append(cell_width)
        cell_width.set(qn("w:type"), "dxa")
        cell_width.set(qn("w:w"), str(width))
        _set_cell_margins(cell)


def _set_cell_margins(cell: object) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (
        ("top", 80),
        ("bottom", 80),
        ("start", 120),
        ("end", 120),
    ):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _format_header_row(row: object) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)
    for cell in row.cells:
        _set_cell_shading(cell, "F2F4F7")
        for run in cell.paragraphs[0].runs:
            run.bold = True


def _set_cell_shading(cell: object, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _fill_row(cells: object, values: object) -> None:
    for cell, value in zip(cells, values, strict=True):
        cell.text = str(value)
    _style_row(cells)


def _style_row(cells: object) -> None:
    row = cells[0]._tc.getparent()
    properties = row.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))
    for cell in cells:
        for paragraph in cell.paragraphs:
            paragraph.style = "GeoFlow Table Text"


def _keep_table_rows_together(table: object) -> None:
    """Keep compact lot tables from splitting into the page header area in LO."""

    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True


def _merge_cells(cells: object, start: int, end: int) -> None:
    cells[start].merge(cells[end])


def _formatted_address(dossier: Dossier) -> str:
    parts: list[str] = []
    if dossier.adresse is not None:
        if dossier.adresse.ligne_voie:
            parts.append(dossier.adresse.ligne_voie)
        if dossier.adresse.complement:
            parts.append(dossier.adresse.complement)
    locality = " ".join(
        item.strip()
        for item in (
            dossier.adresse.code_postal if dossier.adresse else None,
            dossier.commune,
        )
        if item and item.strip()
    )
    if locality:
        parts.append(locality)
    if dossier.departement:
        parts.append(dossier.departement.strip())
    return ", ".join(parts)


def _formatted_cadastre(dossier: Dossier) -> str:
    return "; ".join(
        item.libelle for item in dossier.references_cadastrales if item.libelle
    )


def _zone_labels(dossier: Dossier) -> dict[str, str]:
    labels: dict[str, str] = {}
    for lot in sorted(dossier.lots, key=lambda item: _natural_key(item.numero)):
        for index, zone_id in enumerate(lot.zone_ids, start=1):
            labels[zone_id] = f"Zone {index}"
    common_index = 0
    other_index = 0
    for zone in dossier.zones:
        if zone.id in labels:
            continue
        if zone.categorie == CategorieZone.COMMUNE:
            common_index += 1
            labels[zone.id] = f"Partie commune {common_index}"
        else:
            other_index += 1
            labels[zone.id] = f"Zone technique {other_index}"
    return labels


def _surface(value: float) -> str:
    return f"{value:.2f} m²".replace(".", ",")


def _natural_key(value: str) -> tuple[object, ...]:
    return tuple(
        int(part) if part.isdigit() else part.casefold()
        for part in re.split(r"(\d+)", value)
    )


def _assert_clean_package(content: bytes) -> None:
    try:
        with ZipFile(BytesIO(content)) as archive:
            names = set(archive.namelist())
            if any(
                name.startswith("word/comments")
                or name in {"word/people.xml", "docProps/custom.xml"}
                for name in names
            ):
                raise TemplateCorruptError(
                    "Le document contient des commentaires ou propriétés "
                    "non autorisés."
                )
            for name in names:
                if (
                    name.endswith(".rels")
                    and b'TargetMode="External"' in archive.read(name)
                ):
                    raise TemplateCorruptError(
                        "Le document contient une relation externe."
                    )
            aggregate = b"".join(
                archive.read(name)
                for name in names
                if name.endswith((".xml", ".rels"))
            )
            if re.search(br"<w:(?:ins|del)(?:[ >])", aggregate):
                raise TemplateCorruptError(
                    "Le document contient des marques de révision."
                )
            for marker in (b"file:///", b"C:\\Users\\", b"/home/"):
                if marker in aggregate:
                    raise TemplateCorruptError(
                        "Le document contient un chemin local."
                    )
    except BadZipFile as exc:
        raise TemplateCorruptError(
            "Le document généré n'est pas un DOCX valide."
        ) from exc
